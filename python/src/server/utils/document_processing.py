"""
Document Processing Utilities

This module provides utilities for extracting text from various document formats
including PDF, Word documents, and plain text files.
"""

import asyncio
import io
from collections.abc import Callable
from typing import Any

# Removed direct logging import - using unified config

# Import document processing libraries with availability checks
try:
    import pymupdf4llm

    PYMUPDF4LLM_AVAILABLE = True
except ImportError:
    PYMUPDF4LLM_AVAILABLE = False

try:
    import PyPDF2

    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import pdfplumber

    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from ..config.logfire_config import get_logger, logfire
from .ocr_processing import extract_text_with_ocr, is_ocr_available

logger = get_logger(__name__)


def _preserve_code_blocks_across_pages(text: str) -> str:
    """
    Fix code blocks that were split across PDF page boundaries.
    
    PDFs often break markdown code blocks with page headers like:
    ```python
    def hello():
    --- Page 2 ---
        return "world"
    ```
    
    This function rejoins split code blocks by removing page separators
    that appear within code blocks.
    """
    import re

    # Pattern to match page separators that split code blocks
    # Look for: ``` [content] --- Page N --- [content] ```
    page_break_in_code_pattern = r'(```\w*[^\n]*\n(?:[^`]|`(?!``))*)(\n--- Page \d+ ---\n)((?:[^`]|`(?!``))*)```'

    # Keep merging until no more splits are found
    while True:
        matches = list(re.finditer(page_break_in_code_pattern, text, re.DOTALL))
        if not matches:
            break

        # Replace each match by removing the page separator
        for match in reversed(matches):  # Reverse to maintain positions
            before_page_break = match.group(1)
            page_separator = match.group(2)
            after_page_break = match.group(3)

            # Rejoin the code block without the page separator
            rejoined = f"{before_page_break}\n{after_page_break}```"
            text = text[:match.start()] + rejoined + text[match.end():]

    return text


def _clean_html_to_text(html_content: str) -> str:
    """
    Clean HTML tags and convert to plain text suitable for RAG.
    Preserves code blocks and important structure while removing markup.
    """
    import re

    # First preserve code blocks with their content before general cleaning
    # This ensures code blocks remain intact for extraction
    code_blocks = []

    # Find and temporarily replace code blocks to preserve them
    code_patterns = [
        r'<pre><code[^>]*>(.*?)</code></pre>',
        r'<code[^>]*>(.*?)</code>',
        r'<pre[^>]*>(.*?)</pre>',
    ]

    processed_html = html_content
    placeholder_map = {}

    for pattern in code_patterns:
        matches = list(re.finditer(pattern, processed_html, re.DOTALL | re.IGNORECASE))
        for i, match in enumerate(reversed(matches)):  # Reverse to maintain positions
            # Extract code content and clean HTML entities
            code_content = match.group(1)
            # Clean HTML entities and span tags from code
            code_content = re.sub(r'<span[^>]*>', '', code_content)
            code_content = re.sub(r'</span>', '', code_content)
            code_content = re.sub(r'&lt;', '<', code_content)
            code_content = re.sub(r'&gt;', '>', code_content)
            code_content = re.sub(r'&amp;', '&', code_content)
            code_content = re.sub(r'&quot;', '"', code_content)
            code_content = re.sub(r'&#39;', "'", code_content)

            # Create placeholder
            placeholder = f"__CODE_BLOCK_{len(placeholder_map)}__"
            placeholder_map[placeholder] = code_content.strip()

            # Replace in HTML
            processed_html = processed_html[:match.start()] + placeholder + processed_html[match.end():]

    # Now clean all remaining HTML tags
    # Remove script and style content entirely
    processed_html = re.sub(r'<script[^>]*>.*?</script>', '', processed_html, flags=re.DOTALL | re.IGNORECASE)
    processed_html = re.sub(r'<style[^>]*>.*?</style>', '', processed_html, flags=re.DOTALL | re.IGNORECASE)

    # Convert common HTML elements to readable text
    # Headers
    processed_html = re.sub(r'<h[1-6][^>]*>(.*?)</h[1-6]>', r'\n\n\1\n\n', processed_html, flags=re.DOTALL | re.IGNORECASE)
    # Paragraphs
    processed_html = re.sub(r'<p[^>]*>(.*?)</p>', r'\1\n\n', processed_html, flags=re.DOTALL | re.IGNORECASE)
    # Line breaks
    processed_html = re.sub(r'<br\s*/?>', '\n', processed_html, flags=re.IGNORECASE)
    # List items
    processed_html = re.sub(r'<li[^>]*>(.*?)</li>', r'• \1\n', processed_html, flags=re.DOTALL | re.IGNORECASE)

    # Remove all remaining HTML tags
    processed_html = re.sub(r'<[^>]+>', '', processed_html)

    # Clean up HTML entities
    processed_html = re.sub(r'&nbsp;', ' ', processed_html)
    processed_html = re.sub(r'&lt;', '<', processed_html)
    processed_html = re.sub(r'&gt;', '>', processed_html)
    processed_html = re.sub(r'&amp;', '&', processed_html)
    processed_html = re.sub(r'&quot;', '"', processed_html)
    processed_html = re.sub(r'&#39;', "'", processed_html)
    processed_html = re.sub(r'&#x27;', "'", processed_html)

    # Restore code blocks
    for placeholder, code_content in placeholder_map.items():
        processed_html = processed_html.replace(placeholder, f"\n\n```\n{code_content}\n```\n\n")

    # Clean up excessive whitespace
    processed_html = re.sub(r'\n\s*\n\s*\n', '\n\n', processed_html)  # Max 2 consecutive newlines
    processed_html = re.sub(r'[ \t]+', ' ', processed_html)  # Multiple spaces to single space

    return processed_html.strip()


def extract_text_from_document(file_content: bytes, filename: str, content_type: str) -> str:
    """
    Extract text from various document formats.

    Args:
        file_content: Raw file bytes
        filename: Name of the file
        content_type: MIME type of the file

    Returns:
        Extracted text content

    Raises:
        ValueError: If the file format is not supported
        Exception: If extraction fails
    """
    try:
        # PDF files
        if content_type == "application/pdf" or filename.lower().endswith(".pdf"):
            return extract_text_from_pdf(file_content)

        # Word documents
        elif content_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ] or filename.lower().endswith((".docx", ".doc")):
            return extract_text_from_docx(file_content)

        # HTML files - clean tags and extract text
        elif content_type == "text/html" or filename.lower().endswith((".html", ".htm")):
            # Decode HTML and clean tags for RAG
            html_text = file_content.decode("utf-8", errors="ignore").strip()
            if not html_text:
                raise ValueError(f"The file {filename} appears to be empty.")
            return _clean_html_to_text(html_text)

        # Text files (markdown, txt, etc.)
        elif content_type.startswith("text/") or filename.lower().endswith((
            ".txt",
            ".md",
            ".markdown",
            ".rst",
        )):
            # Decode text and check if it has content
            text = file_content.decode("utf-8", errors="ignore").strip()
            if not text:
                raise ValueError(f"The file {filename} appears to be empty.")
            return text

        else:
            raise ValueError(f"Unsupported file format: {content_type} ({filename})")

    except ValueError:
        # Re-raise ValueError with original message for unsupported formats
        raise
    except Exception as e:
        logfire.error(
            "Document text extraction failed",
            filename=filename,
            content_type=content_type,
            error=str(e),
        )
        # Re-raise with context, preserving original exception chain
        raise Exception(f"Failed to extract text from {filename}") from e


async def extract_text_from_document_async(
    file_content: bytes,
    filename: str,
    content_type: str,
    progress_callback: Callable[[int, int, str], Any] | None = None
) -> str:
    """
    Extract text from various document formats with async progress tracking.

    Args:
        file_content: Raw file bytes
        filename: Name of the file
        content_type: MIME type of the file
        progress_callback: Optional async callback(pages_extracted, total_pages, status_message)

    Returns:
        Extracted text content

    Raises:
        ValueError: If the file format is not supported
        Exception: If extraction fails
    """
    try:
        # PDF files - use async version with progress
        if content_type == "application/pdf" or filename.lower().endswith(".pdf"):
            return await extract_text_from_pdf_async(file_content, progress_callback)

        # Word documents - no page progress (usually fast)
        elif content_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ] or filename.lower().endswith((".docx", ".doc")):
            return extract_text_from_docx(file_content)

        # HTML files - clean tags and extract text
        elif content_type == "text/html" or filename.lower().endswith((".html", ".htm")):
            html_text = file_content.decode("utf-8", errors="ignore").strip()
            if not html_text:
                raise ValueError(f"The file {filename} appears to be empty.")
            return _clean_html_to_text(html_text)

        # Text files (markdown, txt, etc.)
        elif content_type.startswith("text/") or filename.lower().endswith((
            ".txt",
            ".md",
            ".markdown",
            ".rst",
        )):
            text = file_content.decode("utf-8", errors="ignore").strip()
            if not text:
                raise ValueError(f"The file {filename} appears to be empty.")
            return text

        else:
            raise ValueError(f"Unsupported file format: {content_type} ({filename})")

    except ValueError:
        raise
    except Exception as e:
        logfire.error(
            "Document text extraction failed (async)",
            filename=filename,
            content_type=content_type,
            error=str(e),
        )
        raise Exception(f"Failed to extract text from {filename}") from e


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from PDF with Markdown structure preservation.

    Uses pymupdf4llm as primary (best word separation and Markdown output),
    falls back to pdfplumber, then PyPDF2.

    Args:
        file_content: Raw PDF bytes

    Returns:
        Extracted text content (Markdown format when possible)
    """
    if not PYMUPDF4LLM_AVAILABLE and not PDFPLUMBER_AVAILABLE and not PYPDF2_AVAILABLE:
        raise Exception(
            "No PDF processing libraries available. Install pymupdf4llm, pdfplumber, or PyPDF2."
        )

    # Primary: pymupdf4llm (best quality - proper word separation and Markdown)
    if PYMUPDF4LLM_AVAILABLE:
        try:
            import os
            import tempfile

            # pymupdf4llm requires a file path, so write to temp file
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name

            try:
                markdown_text = pymupdf4llm.to_markdown(tmp_path)

                if markdown_text and len(markdown_text.strip()) > 100:
                    logger.info(f"PDF extracted with pymupdf4llm: {len(markdown_text)} chars")
                    return markdown_text
                else:
                    logfire.warning("pymupdf4llm returned insufficient text, trying fallback")
            finally:
                os.unlink(tmp_path)

        except Exception as e:
            logfire.warning(f"pymupdf4llm extraction failed: {e}, trying pdfplumber")

    # Fallback 1: pdfplumber
    if PDFPLUMBER_AVAILABLE:
        try:
            text_content = []
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    except Exception as e:
                        logfire.warning(f"pdfplumber failed on page {page_num + 1}: {e}")
                        continue

            if text_content and len("\n".join(text_content).strip()) > 100:
                combined_text = "\n\n".join(text_content)
                logger.info(f"PDF extracted with pdfplumber: {len(combined_text)} chars")
                return _preserve_code_blocks_across_pages(combined_text)

        except Exception as e:
            logfire.warning(f"pdfplumber extraction failed: {e}, trying PyPDF2")

    # Fallback 2: PyPDF2
    if PYPDF2_AVAILABLE:
        try:
            text_content = []
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    logfire.warning(f"PyPDF2 failed on page {page_num + 1}: {e}")
                    continue

            if text_content:
                combined_text = "\n\n".join(text_content)
                logger.info(f"PDF extracted with PyPDF2: {len(combined_text)} chars")
                return _preserve_code_blocks_across_pages(combined_text)
            # If no text, fall through to OCR

        except Exception as e:
            logfire.warning(f"PyPDF2 extraction failed: {e}, trying OCR")

    # Final fallback: OCR for image-based/scanned PDFs
    if is_ocr_available():
        logger.info("No text extracted - attempting OCR for image-based PDF")
        ocr_text = extract_text_with_ocr(file_content)
        if ocr_text and len(ocr_text.strip()) > 50:
            logger.info(f"PDF extracted with OCR (Tesseract): {len(ocr_text)} chars")
            return ocr_text
        else:
            raise ValueError(
                "No text extracted from PDF: OCR found no readable text. "
                "File may be empty or contain only images without text."
            )
    else:
        raise ValueError(
            "No text extracted from PDF: file appears to be images-only or scanned. "
            "Install OCR dependencies for scanned PDF support: "
            "pip install pytesseract pdf2image (and install tesseract + poppler)"
        )


async def extract_text_from_pdf_async(
    file_content: bytes,
    progress_callback: Callable[[int, int, str], Any] | None = None
) -> str:
    """
    Extract text from PDF with Markdown structure preservation and progress tracking.

    Uses pymupdf4llm with page_chunks=True for page-by-page progress.

    Args:
        file_content: Raw PDF bytes
        progress_callback: Optional async callback(pages_extracted, total_pages, status_message)

    Returns:
        Extracted text content (Markdown format when possible)
    """
    if not PYMUPDF4LLM_AVAILABLE and not PDFPLUMBER_AVAILABLE and not PYPDF2_AVAILABLE:
        raise Exception(
            "No PDF processing libraries available. Install pymupdf4llm, pdfplumber, or PyPDF2."
        )

    async def report_progress(extracted: int, total: int, message: str):
        if progress_callback:
            if asyncio.iscoroutinefunction(progress_callback):
                await progress_callback(extracted, total, message)
            else:
                progress_callback(extracted, total, message)

    # Primary: pymupdf4llm with page_chunks for progress tracking
    if PYMUPDF4LLM_AVAILABLE:
        try:
            import os
            import tempfile

            # pymupdf4llm requires a file path, so write to temp file
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                tmp.write(file_content)
                tmp_path = tmp.name

            try:
                # Use page_chunks=True to get per-page results for progress tracking
                page_data = pymupdf4llm.to_markdown(tmp_path, page_chunks=True)

                if page_data and len(page_data) > 0:
                    total_pages = len(page_data)
                    markdown_parts = []

                    for i, page in enumerate(page_data):
                        page_text = page.get("text", "") if isinstance(page, dict) else str(page)
                        markdown_parts.append(page_text)

                        # Report progress for each page
                        await report_progress(
                            i + 1,
                            total_pages,
                            f"Extracting page {i + 1}/{total_pages}..."
                        )
                        # Yield control to allow progress updates
                        await asyncio.sleep(0)

                    markdown_text = "\n\n".join(markdown_parts)

                    if markdown_text and len(markdown_text.strip()) > 100:
                        logger.info(f"PDF extracted with pymupdf4llm (async): {len(markdown_text)} chars, {total_pages} pages")
                        return markdown_text
                    else:
                        logfire.warning("pymupdf4llm returned insufficient text, trying fallback")
            finally:
                os.unlink(tmp_path)

        except Exception as e:
            logfire.warning(f"pymupdf4llm extraction failed: {e}, trying pdfplumber")

    # Fallback 1: pdfplumber with page progress
    if PDFPLUMBER_AVAILABLE:
        try:
            text_content = []
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                total_pages = len(pdf.pages)
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")

                        await report_progress(
                            page_num + 1,
                            total_pages,
                            f"Extracting page {page_num + 1}/{total_pages}..."
                        )
                        await asyncio.sleep(0)

                    except Exception as e:
                        logfire.warning(f"pdfplumber failed on page {page_num + 1}: {e}")
                        continue

            if text_content and len("\n".join(text_content).strip()) > 100:
                combined_text = "\n\n".join(text_content)
                logger.info(f"PDF extracted with pdfplumber (async): {len(combined_text)} chars")
                return _preserve_code_blocks_across_pages(combined_text)

        except Exception as e:
            logfire.warning(f"pdfplumber extraction failed: {e}, trying PyPDF2")

    # Fallback 2: PyPDF2 with page progress
    if PYPDF2_AVAILABLE:
        try:
            text_content = []
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            total_pages = len(pdf_reader.pages)

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")

                    await report_progress(
                        page_num + 1,
                        total_pages,
                        f"Extracting page {page_num + 1}/{total_pages}..."
                    )
                    await asyncio.sleep(0)

                except Exception as e:
                    logfire.warning(f"PyPDF2 failed on page {page_num + 1}: {e}")
                    continue

            if text_content:
                combined_text = "\n\n".join(text_content)
                logger.info(f"PDF extracted with PyPDF2 (async): {len(combined_text)} chars")
                return _preserve_code_blocks_across_pages(combined_text)

        except Exception as e:
            logfire.warning(f"PyPDF2 extraction failed: {e}, trying OCR")

    # Final fallback: OCR for image-based/scanned PDFs (uses ocr_processing with its own progress)
    from .ocr_processing import extract_text_with_ocr_async, is_ocr_available

    if is_ocr_available():
        logger.info("No text extracted - attempting OCR for image-based PDF")
        ocr_text = await extract_text_with_ocr_async(file_content, progress_callback)
        if ocr_text and len(ocr_text.strip()) > 50:
            logger.info(f"PDF extracted with OCR (async): {len(ocr_text)} chars")
            return ocr_text
        else:
            raise ValueError(
                "No text extracted from PDF: OCR found no readable text. "
                "File may be empty or contain only images without text."
            )
    else:
        raise ValueError(
            "No text extracted from PDF: file appears to be images-only or scanned. "
            "Install OCR dependencies for scanned PDF support: "
            "pip install pytesseract pdf2image (and install tesseract + poppler)"
        )


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from Word documents (.docx).

    Args:
        file_content: Raw DOCX bytes

    Returns:
        Extracted text content
    """
    if not DOCX_AVAILABLE:
        raise Exception("python-docx library not available. Please install python-docx.")

    try:
        doc = DocxDocument(io.BytesIO(file_content))
        text_content = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))

        if not text_content:
            raise ValueError("No text content found in document")

        return "\n\n".join(text_content)

    except Exception as e:
        raise Exception("Failed to extract text from Word document") from e
